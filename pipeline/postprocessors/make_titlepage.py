#!/usr/bin/env python3
"""
Build university titlepage as a .docx.
Output: titlepage.docx in same directory.
"""
import os
from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.join(os.path.dirname(__file__), "titlepage.docx")

doc = Document()

# page setup
section = doc.sections[0]
section.top_margin = Mm(20)
section.bottom_margin = Mm(20)
section.left_margin = Mm(30)
section.right_margin = Mm(15)
section.page_height = Mm(297)
section.page_width = Mm(210)

# Normal font
styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(14)
rpr = normal.element.get_or_add_rPr()
rf = rpr.find(qn("w:rFonts"))
if rf is None:
    rf = OxmlElement("w:rFonts")
    rpr.append(rf)
for k in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
    rf.set(qn(k), "Times New Roman")

def add_para(text, align=WD_ALIGN_PARAGRAPH.CENTER, bold=False, size=14,
             space_before=0, space_after=0, line_spacing=1.15, indent=0):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    pf.first_line_indent = Cm(indent)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    # force font explicitly
    r_rpr = run._element.get_or_add_rPr()
    rf2 = OxmlElement("w:rFonts")
    for k in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf2.set(qn(k), "Times New Roman")
    r_rpr.append(rf2)
    return p

def add_runs(align, *runs, space_before=0, space_after=0, line_spacing=1.15, indent=0):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    pf.first_line_indent = Cm(indent)
    for text, bold, size in runs:
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(size)
        r.bold = bold
        r_rpr = r._element.get_or_add_rPr()
        rf2 = OxmlElement("w:rFonts")
        for k in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rf2.set(qn(k), "Times New Roman")
        r_rpr.append(rf2)
    return p

# Remove the default empty paragraph
body = doc.element.body
default_p = body.find(qn("w:p"))
if default_p is not None:
    body.remove(default_p)

# --- Header block (centered) ---
add_para("Министерство науки и высшего образования Российской Федерации",
         bold=True, size=14, line_spacing=1.15, space_after=2)
add_para("ФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ АВТОНОМНОЕ ОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ ВЫСШЕГО ОБРАЗОВАНИЯ",
         bold=False, size=10, line_spacing=1.15, space_after=2)
add_para("«Название университета»",
         bold=True, size=14, line_spacing=1.15, space_after=2)
add_para("(Университет)",
         bold=False, size=14, line_spacing=1.15, space_after=12)

# --- Metadata block (left-aligned) ---
add_runs(WD_ALIGN_PARAGRAPH.LEFT,
         ("Факультет ", True, 14),
         ("название факультета", False, 14),
         space_before=12, line_spacing=1.3)
add_runs(WD_ALIGN_PARAGRAPH.LEFT,
         ("Образовательная программа ", True, 14),
         ("Название образовательной программы", False, 14),
         line_spacing=1.3)
add_runs(WD_ALIGN_PARAGRAPH.LEFT,
         ("Направление подготовки ", True, 14),
         ("XX.XX.XX Название направления", False, 14),
         line_spacing=1.3)
add_runs(WD_ALIGN_PARAGRAPH.LEFT,
         ("Квалификация ", True, 14),
         ("Бакалавр", False, 14),
         line_spacing=1.3, space_after=24)

# --- Title block (centered) — финальный ВКР, НЕ «отчёт проектного семинара» ---
add_para("ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА",
         bold=True, size=14, line_spacing=1.3, space_before=24)
add_para("(бакалаврская работа)",
         bold=False, size=14, line_spacing=1.3, space_after=6)

add_runs(WD_ALIGN_PARAGRAPH.LEFT,
         ("Тема ", True, 14),
         ("«Тема выпускной квалификационной работы»", False, 14),
         space_before=12, line_spacing=1.3, space_after=12)

# --- Spacer then right-aligned author/supervisor ---
for _ in range(4):
    add_para("", line_spacing=1.0, space_after=0)

add_para("Выполнил:", align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True, size=14, line_spacing=1.3)
add_para("Фамилия Имя Отчество, Группа № XXXXX", align=WD_ALIGN_PARAGRAPH.RIGHT,
         bold=False, size=14, line_spacing=1.3, space_after=2)
add_para("Научный руководитель:", align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True, size=14, line_spacing=1.3)
add_para("Фамилия Имя Отчество, преподаватель факультета название факультета", align=WD_ALIGN_PARAGRAPH.RIGHT,
         bold=False, size=14, line_spacing=1.3, space_after=12)

# --- Footer: city + year centered at bottom ---
for _ in range(2):
    add_para("", line_spacing=1.0, space_after=0)
add_para("Санкт-Петербург", align=WD_ALIGN_PARAGRAPH.CENTER, bold=False, size=14, line_spacing=1.15)
add_para("2026", align=WD_ALIGN_PARAGRAPH.CENTER, bold=False, size=14, line_spacing=1.15)

# Page break at the end so body starts on new page
p = doc.add_paragraph()
r = p.add_run()
r.add_break(WD_BREAK.PAGE)

doc.save(OUT)
print(f"Saved: {OUT}")
