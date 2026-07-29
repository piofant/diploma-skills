#!/usr/bin/env python3
"""
Generate reference.docx for pandoc: GOST + university styles.
Times New Roman 14pt, margins 30/15/20/20mm, 1.5 spacing, 1.25cm indent.
Forces black color + Times New Roman on all headings (override theme defaults).
"""
import os
from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.join(os.path.dirname(__file__), "template.docx")


def ensure_rfonts(rpr, font="Times New Roman"):
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.insert(0, rf)
    for k in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(k), font)


def ensure_color(rpr, hex_color="000000"):
    c = rpr.find(qn("w:color"))
    if c is None:
        c = OxmlElement("w:color")
        rpr.append(c)
    c.set(qn("w:val"), hex_color)
    c.set(qn("w:themeColor"), "none") if False else None
    # remove theme color attrs if any
    for attr in ("w:themeColor", "w:themeTint", "w:themeShade"):
        if c.get(qn(attr)) is not None:
            del c.attrib[qn(attr)]


def ensure_size(rpr, pt_size):
    sz = rpr.find(qn("w:sz"))
    if sz is None:
        sz = OxmlElement("w:sz")
        rpr.append(sz)
    sz.set(qn("w:val"), str(int(pt_size * 2)))
    szcs = rpr.find(qn("w:szCs"))
    if szcs is None:
        szcs = OxmlElement("w:szCs")
        rpr.append(szcs)
    szcs.set(qn("w:val"), str(int(pt_size * 2)))


def ensure_lang(rpr):
    el = rpr.find(qn("w:lang"))
    if el is None:
        el = OxmlElement("w:lang")
        rpr.append(el)
    el.set(qn("w:val"), "ru-RU")
    el.set(qn("w:eastAsia"), "ru-RU")


def set_bold(rpr, bold=True):
    b = rpr.find(qn("w:b"))
    if bold:
        if b is None:
            b = OxmlElement("w:b")
            rpr.append(b)
    else:
        if b is not None:
            rpr.remove(b)


def set_italic(rpr, italic=False):
    i = rpr.find(qn("w:i"))
    if italic:
        if i is None:
            i = OxmlElement("w:i")
            rpr.append(i)
    else:
        if i is not None:
            rpr.remove(i)
    # Force-disable via explicit val="false"
    if not italic and i is None:
        i = OxmlElement("w:i")
        i.set(qn("w:val"), "false")
        rpr.append(i)


doc = Document()

# ---- Page setup ----
section = doc.sections[0]
section.top_margin = Mm(20)
section.bottom_margin = Mm(20)
section.left_margin = Mm(30)
section.right_margin = Mm(15)
section.page_height = Mm(297)
section.page_width = Mm(210)

styles = doc.styles


def configure_style(name, base=None, size=14, bold=False, color="000000",
                    align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=None,
                    left_indent=None, space_before=0, space_after=0,
                    line_spacing=1.5, keep_next=False):
    try:
        st = styles[name]
    except KeyError:
        st = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    if base is not None:
        try:
            st.base_style = styles[base]
        except KeyError:
            pass
    # run properties
    rpr = st.element.get_or_add_rPr()
    ensure_rfonts(rpr)
    ensure_size(rpr, size)
    ensure_color(rpr, color)
    ensure_lang(rpr)
    set_bold(rpr, bold)
    set_italic(rpr, italic=False)
    # paragraph properties
    pf = st.paragraph_format
    pf.alignment = align
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    if left_indent is not None:
        pf.left_indent = left_indent
    pf.keep_with_next = keep_next
    return st


# Normal — базовый
configure_style("Normal", size=14, bold=False, color="000000",
                align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                first_line_indent=Cm(1.25), line_spacing=1.5,
                space_before=0, space_after=0)

# Body Text + First Paragraph — те же параметры
configure_style("Body Text", base="Normal", size=14, bold=False, color="000000",
                align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                first_line_indent=Cm(1.25), line_spacing=1.5)

configure_style("First Paragraph", base="Normal", size=14, bold=False, color="000000",
                align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                first_line_indent=Cm(1.25), line_spacing=1.5)

configure_style("Compact", base="Normal", size=14, bold=False, color="000000",
                align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                first_line_indent=Cm(1.25), line_spacing=1.5)

# Headings — жирные, чёрные, TNR
# вуза ВКР 2026 регламент: «Заголовки разделов и подразделов пишутся с абзацного
# отступа (1,25 см)... Дополнительных интервалов до и после заголовков быть
# не должно».  → first_line_indent=1.25 cm, space_before/after=0.
# Структурные элементы (ВВЕДЕНИЕ, ЗАКЛЮЧЕНИЕ, ...) ловит fix_structural_elements.py
# и перевыравнивает по центру с indent=0.
configure_style("Heading 1", size=14, bold=True, color="000000",
                align=WD_ALIGN_PARAGRAPH.LEFT,
                first_line_indent=Cm(1.25), space_before=0, space_after=0,
                line_spacing=1.5, keep_next=True)

configure_style("Heading 2", size=14, bold=True, color="000000",
                align=WD_ALIGN_PARAGRAPH.LEFT,
                first_line_indent=Cm(1.25), space_before=0, space_after=0,
                line_spacing=1.5, keep_next=True)

configure_style("Heading 3", size=14, bold=True, color="000000",
                align=WD_ALIGN_PARAGRAPH.LEFT,
                first_line_indent=Cm(1.25), space_before=0, space_after=0,
                line_spacing=1.5, keep_next=True)

# (Heading 4 переопределён ниже на НЕ bold — требование секретаря ГЭК)

# Heading 4-9 — НЕ выделяются полужирным (требование секретаря ГЭК секретаря ГЭК, 2026-05-14):
# «Третий и далее уровень нумерованного заголовка не выделяется полужирным шрифтом»
# В нашем mapping: «## 1 АНАЛИЗ» → H2 (1й ур), «### 1.1» → H3 (2й), «#### 1.1.1» → H4 (3й ур) → НЕ bold
configure_style("Heading 4", size=14, bold=False, color="000000",
                align=WD_ALIGN_PARAGRAPH.LEFT,
                first_line_indent=Cm(1.25), space_before=0, space_after=0,
                line_spacing=1.5, keep_next=True)

for hn in (5, 6, 7, 8, 9):
    configure_style(f"Heading {hn}", size=14, bold=False, color="000000",
                    align=WD_ALIGN_PARAGRAPH.LEFT,
                    first_line_indent=Cm(1.25), space_before=0, space_after=0,
                    line_spacing=1.5, keep_next=True)

# TOC styles
configure_style("TOC Heading", size=14, bold=True, color="000000",
                align=WD_ALIGN_PARAGRAPH.CENTER,
                first_line_indent=Cm(0), space_before=0, space_after=18,
                line_spacing=1.5)

configure_style("toc 1", base="Normal", size=14, bold=False, color="000000",
                align=WD_ALIGN_PARAGRAPH.LEFT,
                first_line_indent=Cm(0), left_indent=Cm(0),
                line_spacing=1.5, space_before=0, space_after=0)

configure_style("toc 2", base="Normal", size=14, bold=False, color="000000",
                align=WD_ALIGN_PARAGRAPH.LEFT,
                first_line_indent=Cm(0), left_indent=Cm(0.75),
                line_spacing=1.5, space_before=0, space_after=0)

configure_style("toc 3", base="Normal", size=14, bold=False, color="000000",
                align=WD_ALIGN_PARAGRAPH.LEFT,
                first_line_indent=Cm(0), left_indent=Cm(1.5),
                line_spacing=1.5, space_before=0, space_after=0)

# List Paragraph (для нумерованных списков от pandoc)
configure_style("List Paragraph", base="Normal", size=14, bold=False, color="000000",
                align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                first_line_indent=Cm(0), left_indent=Cm(1.25),
                line_spacing=1.5, space_before=0, space_after=0)

# ---- Footer: page number centered (arabic) ----
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
# clear existing runs
for r in list(fp.runs):
    r._element.getparent().remove(r._element)
run = fp.add_run()
run.font.name = "Times New Roman"
run.font.size = Pt(14)
r_rpr = run._element.get_or_add_rPr()
ensure_rfonts(r_rpr)
ensure_size(r_rpr, 14)
ensure_color(r_rpr, "000000")
fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
instr.text = " PAGE \\* MERGEFORMAT "
fld_sep = OxmlElement("w:fldChar"); fld_sep.set(qn("w:fldCharType"), "separate")
fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
run._r.append(fld_begin); run._r.append(instr); run._r.append(fld_sep); run._r.append(fld_end)

doc.save(OUT)
print(f"Saved: {OUT}")
