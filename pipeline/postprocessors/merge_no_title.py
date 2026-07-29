#!/usr/bin/env python3
"""
Merge WITHOUT titlepage — для заливки в ИСУ (секретарь ГЭК отвергает PDF
с собственным титулом, ИСУ генерирует его автоматически из формы).

В остальном идентичен merge.py: добавляет «СОДЕРЖАНИЕ» + TOC-поле перед
первым параграфом тела + разрыв страницы.
"""
import os
import sys
from docx import Document
from docx.shared import Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(__file__)
BODY = os.path.join(HERE, "body.docx")
OUT = os.path.join(HERE, "diploma_body.docx")


def insert_toc_field():
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), "TOCHeading")
    pPr.append(pStyle)
    p.append(pPr)
    run_begin = OxmlElement("w:r")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run_begin.append(fld_begin)
    p.append(run_begin)
    run_instr = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r' TOC \o "1-2" \h \z \u '
    run_instr.append(instr)
    p.append(run_instr)
    run_sep = OxmlElement("w:r")
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run_sep.append(fld_sep)
    p.append(run_sep)
    run_end = OxmlElement("w:r")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_end.append(fld_end)
    p.append(run_end)
    return p


def fix_paragraph_indent_for_headings(doc):
    for para in doc.paragraphs:
        st = para.style.name if para.style else ""
        if st.startswith("Heading"):
            para.paragraph_format.first_line_indent = Cm(0)


def force_font(doc, font="Times New Roman"):
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = font
            rpr = run._element.get_or_add_rPr()
            rf = rpr.find(qn("w:rFonts"))
            if rf is None:
                rf = OxmlElement("w:rFonts")
                rpr.append(rf)
            for k in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
                rf.set(qn(k), font)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = font
                        rpr = run._element.get_or_add_rPr()
                        rf = rpr.find(qn("w:rFonts"))
                        if rf is None:
                            rf = OxmlElement("w:rFonts")
                            rpr.append(rf)
                        for k in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
                            rf.set(qn(k), font)


def main():
    body = Document(BODY)
    fix_paragraph_indent_for_headings(body)
    force_font(body)

    body_el = body.element.body
    first_para = body_el.find(qn("w:p"))
    toc_para = insert_toc_field()

    toc_head = OxmlElement("w:p")
    toc_head_pPr = OxmlElement("w:pPr")
    toc_head_pStyle = OxmlElement("w:pStyle")
    toc_head_pStyle.set(qn("w:val"), "TOCHeading")
    toc_head_pPr.append(toc_head_pStyle)
    toc_head.append(toc_head_pPr)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    b = OxmlElement("w:b")
    rPr.append(b)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "28")
    rPr.append(sz)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = "СОДЕРЖАНИЕ"
    r.append(t)
    toc_head.append(r)

    pb_p = OxmlElement("w:p")
    pb_r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    pb_r.append(br)
    pb_p.append(pb_r)

    if first_para is not None:
        first_para.addprevious(toc_head)
        first_para.addprevious(toc_para)
        first_para.addprevious(pb_p)

    body.save(OUT)
    print(f"Saved (no title): {OUT}")


if __name__ == "__main__":
    main()
